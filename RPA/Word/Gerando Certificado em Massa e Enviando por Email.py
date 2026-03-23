from docx import Document as Doc
from docx.shared import Pt
from openpyxl import load_workbook as lw
from docx.shared import RGBColor
import win32com.client as win32

outlook = win32.Dispatch("Outlook.Application")

nome_arquivo_alunos = "F:\\Python\\Word\\DadosAlunos.xlsx" #por estar na mesma pasta do código, não é necessário pegar caminho

planilhaDadosAlunos = lw(nome_arquivo_alunos)

#Selecionando a sheet/planilha/aba
sheet_selecionada = planilhaDadosAlunos['Nomes']

for linha in range(2, len(sheet_selecionada['A'])+1):
    #Abre o arquivo word
    arquivoWord = Doc("F:\\Python\\Word\\Certificado3.docx") #Pegar o template depois

    #Seleciona o estilo
    estilo = arquivoWord.styles["Normal"]

    #Pegando os dados do aluno quando passar na célula
    nomeAluno = sheet_selecionada['A%s' % linha].value
    dia = sheet_selecionada['B%s' % linha].value
    mes = sheet_selecionada['C%s' % linha].value
    ano = sheet_selecionada['D%s' % linha].value
    curso = sheet_selecionada['E%s' % linha].value
    instrutor = sheet_selecionada['F%s' % linha].value
    emailAluno = sheet_selecionada['G%s' % linha].value

    for paragraph in arquivoWord.paragraphs:
        if "@nome" in paragraph.text:
            paragraph.text = nomeAluno
            fonte = estilo.font
            fonte.name = "Calibri (Corpo)"
            fonte.size = Pt(24)


        paragraph_1 = "Concluiu com sucesso o curso de "
        paragraph_2 = ", com carga horária de 20 horas, promovido pela escola de Cursos Online em "
        paragraphComplement = f'{paragraph_2} {dia} de {mes} de {ano}'
        if "escola" in paragraph.text:
            paragraph.text = paragraph_1
            fonte = estilo.font
            fonte.name = "Calibri (Corpo)"
            fonte.size = Pt(24)
            addNewWord = paragraph.add_run(curso)
            addNewWord.font.color.rgb = RGBColor(255,0,0)
            addNewWord.underline = True
            addNewWord.bold = True

            addNewWord = paragraph.add_run(paragraphComplement)
            addNewWord.font.color.rgb = RGBColor(0, 0, 0)

        if "Instrutor" in paragraph.text:
            paragraph.text = instrutor + "- Instrutor"
            fonte = estilo.font
            fonte.name = "Calibri (Corpo)"
            fonte.size = Pt(24)

    #Pegando o caminho da pasta e configurando o nome do certificado
    caminhoCertificados = "F:\\Python\\Word\\Certificados\\"+nomeAluno+".docx"

    #Salvando o certificado do aluno
    arquivoWord.save(caminhoCertificados)

    emailOutlook = outlook.CreateItem(0)
    emailOutlook.To = emailAluno
    emailOutlook.Subject = "Certificado "+nomeAluno
    emailOutlook.HTMLBody = f""" 
        <p>Boa noite {nomeAluno}</p>
        <p>Segue seu <b>certificado</b></p>
        <p>Atenciosamente</p>
    """

    emailOutlook.Attachments.Add(caminhoCertificados)

    #Salvar como Draft no outlook
    emailOutlook.save()

print("Certificados gerados com sucesso!")